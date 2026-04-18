"""
验房报告生成模块
生成符合规范的施工质量报告
"""

import os
from datetime import datetime
from typing import Dict, List, Optional
from dataclasses import dataclass
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@dataclass
class InspectionItem:
    """检测项"""
    id: int
    name: str
    standard: str  # 规范值/允许误差
    bim_value: Optional[float]  # BIM模型设计值
    measured_value: Optional[float]  # 点云实测值
    deviation: Optional[float]  # 偏差
    is_passed: Optional[bool]  # 是否合格


# 标准检测项定义
INSPECTION_ITEMS = [
    {"id": 1, "name": "室内净高", "standard": "偏差≤5mm", "tolerance": 5},
    {"id": 2, "name": "长", "standard": "误差≤±10mm", "tolerance": 10},
    {"id": 3, "name": "宽", "standard": "误差≤±10mm", "tolerance": 10},
    {"id": 4, "name": "墙面垂直度", "standard": "误差≤3mm", "tolerance": 3},
    {"id": 5, "name": "地面水平度", "standard": "误差≤±3mm", "tolerance": 3},
    {"id": 6, "name": "阴阳角方正", "standard": "误差≤2mm", "tolerance": 2},
    {"id": 7, "name": "门洞尺寸", "standard": "宽×高，误差≤±3mm", "tolerance": 3},
    {"id": 8, "name": "窗户洞口尺寸", "standard": "宽×高，误差≤±3mm", "tolerance": 3},
]


class ReportGenerator:
    """验房报告生成器"""

    def __init__(self, room_name: str = "Xxx房间"):
        self.room_name = room_name
        self.items: List[InspectionItem] = []
        self.created_time = datetime.now().strftime("%Y-%m-%d %H:%M")

    def add_result(self,
                   item_id: int,
                   bim_value: float,
                   measured_value: float,
                   tolerance_override: Optional[float] = None) -> InspectionItem:
        """
        添加检测结果

        Args:
            item_id: 检测项编号 (1-8)
            bim_value: BIM设计值
            measured_value: 点云实测值
            tolerance_override: 自定义容差（可选）

        Returns:
            InspectionItem: 检测项结果
        """
        # 获取标准信息
        item_def = INSPECTION_ITEMS[item_id - 1] if 1 <= item_id <= 8 else None
        if item_def is None:
            raise ValueError(f"无效的检测项编号: {item_id}")

        tolerance = tolerance_override or item_def["tolerance"]

        # 计算偏差
        deviation = measured_value - bim_value

        # 判断是否合格
        is_passed = abs(deviation) <= tolerance

        item = InspectionItem(
            id=item_id,
            name=item_def["name"],
            standard=item_def["standard"],
            bim_value=bim_value,
            measured_value=measured_value,
            deviation=deviation,
            is_passed=is_passed
        )

        self.items.append(item)
        return item

    def set_room_height(self, bim_height: float, measured_height: float):
        """设置房间净高"""
        return self.add_result(1, bim_height, measured_height)

    def set_room_length(self, bim_length: float, measured_length: float):
        """设置房间长度"""
        return self.add_result(2, bim_length, measured_length)

    def set_room_width(self, bim_width: float, measured_width: float):
        """设置房间宽度"""
        return self.add_result(3, bim_width, measured_width)

    def set_wall_verticality(self, bim_vertical: float, measured_vertical: float):
        """设置墙面垂直度"""
        return self.add_result(4, bim_vertical, measured_vertical)

    def set_floor_levelness(self, bim_level: float, measured_level: float):
        """设置地面水平度"""
        return self.add_result(5, bim_level, measured_level)

    def set_corner_squareness(self, bim_corner: float, measured_corner: float):
        """设置阴阳角方正"""
        return self.add_result(6, bim_corner, measured_corner)

    def set_door_size(self, bim_width: float, bim_height: float,
                      measured_width: float, measured_height: float):
        """设置门洞尺寸"""
        # 门洞需要检查宽和高两个值
        self.add_result(7, bim_width, measured_width)
        # 可以单独记录高度偏差

    def set_window_size(self, bim_width: float, bim_height: float,
                        measured_width: float, measured_height: float):
        """设置窗户洞口尺寸"""
        self.add_result(8, bim_width, measured_width)

    def generate_dataframe(self) -> pd.DataFrame:
        """生成数据表格"""
        data = []
        for item in self.items:
            row = {
                "检测项": item.id,
                "项目名称": item.name,
                "规范值": item.standard,
                "BIM模型设计值": self._format_value(item.bim_value),
                "点云模型实测值": self._format_value(item.measured_value),
                "偏差": self._format_deviation(item.deviation),
                "是否合格": self._format_status(item.is_passed)
            }
            data.append(row)

        return pd.DataFrame(data)

    def _format_value(self, value: Optional[float]) -> str:
        """格式化数值"""
        if value is None:
            return ""
        return f"{value:.2f} mm" if abs(value) < 1000 else f"{value/1000:.3f} m"

    def _format_deviation(self, deviation: Optional[float]) -> str:
        """格式化偏差"""
        if deviation is None:
            return ""
        sign = "+" if deviation >= 0 else ""
        return f"{sign}{deviation:.2f} mm"

    def _format_status(self, is_passed: Optional[bool]) -> str:
        """格式化合格状态"""
        if is_passed is None:
            return ""
        return "合格" if is_passed else "不合格"

    def export_excel(self, filepath: str) -> bool:
        """
        导出为 Excel 格式报告

        Args:
            filepath: 输出文件路径

        Returns:
            bool: 是否成功
        """
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "施工质量报告"

            # 设置标题
            ws.merge_cells('A1:G1')
            ws['A1'] = f"{self.room_name}施工质量报告"
            ws['A1'].font = Font(size=16, bold=True)
            ws['A1'].alignment = Alignment(horizontal='center')

            # 设置时间
            ws['A2'] = f"检测时间: {self.created_time}"
            ws['A2'].alignment = Alignment(horizontal='left')

            # 表头
            headers = ["检测项", "项目名称", "规范值", "BIM模型设计值",
                       "点云模型实测值", "偏差", "是否合格"]
            header_row = 4

            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=header_row, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')

            # 设置边框样式
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # 数据行
            for row_idx, item in enumerate(self.items, header_row + 1):
                ws.cell(row=row_idx, column=1, value=item.id)
                ws.cell(row=row_idx, column=2, value=item.name)
                ws.cell(row=row_idx, column=3, value=item.standard)
                ws.cell(row=row_idx, column=4, value=self._format_value(item.bim_value))
                ws.cell(row=row_idx, column=5, value=self._format_value(item.measured_value))
                ws.cell(row=row_idx, column=6, value=self._format_deviation(item.deviation))

                status_cell = ws.cell(row=row_idx, column=7, value=self._format_status(item.is_passed))
                if item.is_passed:
                    status_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                    status_cell.font = Font(color="006100")
                else:
                    status_cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                    status_cell.font = Font(color="9C0006")

                # 设置边框和对齐
                for col in range(1, 8):
                    cell = ws.cell(row=row_idx, column=col)
                    cell.border = thin_border
                    cell.alignment = Alignment(horizontal='center')

            # 设置表头边框
            for col in range(1, 8):
                ws.cell(row=header_row, column=col).border = thin_border

            # 调整列宽
            ws.column_dimensions['A'].width = 10
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 20
            ws.column_dimensions['D'].width = 18
            ws.column_dimensions['E'].width = 18
            ws.column_dimensions['F'].width = 12
            ws.column_dimensions['G'].width = 10

            wb.save(filepath)
            print(f"✓ 已导出报告: {filepath}")
            return True

        except Exception as e:
            print(f"✗ 导出失败: {e}")
            return False

    def print_report(self):
        """打印报告"""
        print(f"\n{'='*60}")
        print(f"{self.room_name}施工质量报告")
        print(f"检测时间: {self.created_time}")
        print(f"{'='*60}")

        print(f"\n{'检测项':<6}{'项目名称':<12}{'规范值':<20}{'BIM设计值':<15}{'实测值':<15}{'偏差':<10}{'是否合格'}")
        print("-" * 90)

        for item in self.items:
            bim_val = self._format_value(item.bim_value)
            meas_val = self._format_value(item.measured_value)
            dev = self._format_deviation(item.deviation)
            status = self._format_status(item.is_passed)

            print(f"{item.id:<6}{item.name:<12}{item.standard:<20}{bim_val:<15}{meas_val:<15}{dev:<10}{status}")

        print(f"\n{'='*60}")

        # 统计
        passed = sum(1 for i in self.items if i.is_passed)
        failed = sum(1 for i in self.items if i.is_passed is False)
        print(f"汇总: 合格 {passed}项, 不合格 {failed}项")

    def export_word(self, filepath: str) -> bool:
        """
        导出为 Word 格式报告

        Args:
            filepath: 输出文件路径

        Returns:
            bool: 是否成功
        """
        try:
            doc = Document()

            # 设置中文字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 标题
            title = doc.add_heading(f"{self.room_name}施工质量报告", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 检测时间
            time_para = doc.add_paragraph()
            time_para.add_run(f"检测时间: {self.created_time}")
            time_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 空行
            doc.add_paragraph()

            # 创建表格
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头
            headers = ["检测项", "项目名称", "规范值", "BIM模型设计值", "点云模型实测值", "偏差", "是否合格"]
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # 设置表头格式
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            # 数据行
            for item in self.items:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.id)
                row_cells[1].text = item.name
                row_cells[2].text = item.standard
                row_cells[3].text = self._format_value(item.bim_value)
                row_cells[4].text = self._format_value(item.measured_value)
                row_cells[5].text = self._format_deviation(item.deviation)
                row_cells[6].text = self._format_status(item.is_passed)

                # 设置单元格格式
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

                # 设置是否合格单元格颜色
                status_cell = row_cells[6]
                status_para = status_cell.paragraphs[0]
                if item.is_passed:
                    # 绿色
                    for run in status_para.runs:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    # 红色
                    for run in status_para.runs:
                        run.font.color.rgb = RGBColor(192, 0, 0)

            # 空行
            doc.add_paragraph()

            # 汇总
            passed = sum(1 for i in self.items if i.is_passed)
            failed = sum(1 for i in self.items if i.is_passed is False)
            summary = doc.add_paragraph()
            summary.add_run(f"汇总: 合格 {passed}项, 不合格 {failed}项").bold = True

            # 保存
            doc.save(filepath)
            print(f"✓ 已导出Word报告: {filepath}")
            return True

        except Exception as e:
            print(f"✗ 导出Word失败: {e}")
            return False


# 测试
if __name__ == "__main__":
    report = ReportGenerator("教学楼101室")

    # 模拟数据
    report.set_room_height(3000, 3005)  # 偏差 5mm，刚好合格
    report.set_room_length(5500, 5512)  # 偏差 12mm，不合格
    report.set_room_width(4000, 4003)   # 偏差 3mm，合格
    report.set_wall_verticality(0, 2.5)  # 偏差 2.5mm，合格
    report.set_floor_levelness(0, 4)    # 偏差 4mm，不合格

    report.print_report()
    report.export_excel("/tmp/test_report.xlsx")
    report.export_word("/tmp/test_report.docx")